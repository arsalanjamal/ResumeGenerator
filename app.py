import streamlit as st
from transformers import pipeline
from fpdf import FPDF
from io import BytesIO
from docx import Document  # New iimport streamlit as st
from transformers import pipeline
from fpdf import FPDF
from io import BytesIO
from docx import Document  # For Word format support
import re

# Load pre-trained model for resume generation
pipe_resume = pipeline("text2text-generation", model="nakamoto-yama/t5-resume-generation")

# Function to extract keywords from the job description
def extract_keywords(job_description):
    keywords = set(re.findall(r'\b\w+\b', job_description.lower()))
    return keywords

# Function to generate the resume
def generate_resume(name, job_role, education, skills, experience, job_description):
    keywords = extract_keywords(job_description)
    input_text = f"Generate a professional, ATS-optimized resume for a {job_role}. The person’s name is {name}, with education in {education}. The person has the following skills: {skills}. The experience includes: {experience}. Focus on the following keywords from the job description: {', '.join(keywords)}."
    resume = pipe_resume(input_text)[0]['generated_text']
    return resume

# Function to export the resume to PDF with a background color
def export_to_pdf(name, job_role, resume_text, education, skills, experience, phone, email, linkedin, address, background_color="white"):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_fill_color(255, 255, 255)  # Default white

    # Set background color based on selection
    if background_color == "light grey":
        pdf.set_fill_color(240, 240, 240)
    elif background_color == "blue":
        pdf.set_fill_color(230, 240, 255)

    pdf.rect(0, 0, 210, 297, 'F')  # Apply background color

    # Set font for the resume
    pdf.set_font("Arial", size=16, style='B')
    
    # Add Title (capitalized)
    pdf.cell(200, 10, txt=name.upper(), ln=True, align="C")
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"Job Role: {job_role.upper()}", ln=True, align="C")
    pdf.ln(10)

    # Add Contact Info
    pdf.set_font("Arial", size=10)
    pdf.cell(200, 10, txt=f"Phone: {phone} | Email: {email} | LinkedIn: {linkedin} | Address: {address}", ln=True, align="C")
    pdf.ln(10)

    # Add Professional Summary Section
    pdf.set_font("Arial", style='B', size=12)
    pdf.cell(200, 10, txt="PROFESSIONAL SUMMARY", ln=True, align="L")
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=resume_text)
    pdf.ln(10)

    # Add Education Section
    pdf.set_font("Arial", style='B', size=12)
    pdf.cell(200, 10, txt="EDUCATION", ln=True, align="L")
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=education)
    pdf.ln(10)

    # Add Skills Section
    pdf.set_font("Arial", style='B', size=12)
    pdf.cell(200, 10, txt="SKILLS", ln
mport for DOCX support
import re

# Load pre-trained model for resume generation
pipe_resume = pipeline("text2text-generation", model="nakamoto-yama/t5-resume-generation")

# Function to extract keywords from the job description
def extract_keywords(job_description):
    keywords = set(re.findall(r'\b\w+\b', job_description.lower()))
    return keywords

# Function to generate the resume
def generate_resume(name, job_role, education, skills, experience, job_description):
    keywords = extract_keywords(job_description)
    input_text = f"Generate a professional, ATS-optimized resume for a {job_role}. The person’s name is {name}, with education in {education}. The person has the following skills: {skills}. The experience includes: {experience}. Focus on the following keywords from the job description: {', '.join(keywords)}."
    resume = pipe_resume(input_text)[0]['generated_text']
    return resume

# Function to export the resume to PDF with a background color
def export_to_pdf(name, job_role, resume_text, education, skills, experience, phone, email, linkedin, address, background_color="white"):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_fill_color(255, 255, 255)  # Default white

    # Set background color based on selection
    if background_color == "light grey":
        pdf.set_fill_color(240, 240, 240)
    elif background_color == "blue":
        pdf.set_fill_color(230, 240, 255)

    pdf.rect(0, 0, 210, 297, 'F')  # Apply background color

    # Set font for the resume
    pdf.set_font("Arial", size=16, style='B')
    
    # Add Title (capitalized)
    pdf.cell(200, 10, txt=name.upper(), ln=True, align="C")
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"Job Role: {job_role.upper()}", ln=True, align="C")
    pdf.ln(10)

    # Add Contact Info
    pdf.set_font("Arial", size=10)
    pdf.cell(200, 10, txt=f"Phone: {phone} | Email: {email} | LinkedIn: {linkedin} | Address: {address}", ln=True, align="C")
    pdf.ln(10)

    # Add Professional Summary Section
    pdf.set_font("Arial", style='B', size=12)
    pdf.cell(200, 10, txt="PROFESSIONAL SUMMARY", ln=True, align="L")
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=resume_text)
    pdf.ln(10)

    # Add Education Section
    pdf.set_font("Arial", style='B', size=12)
    pdf.cell(200, 10, txt="EDUCATION", ln=True, align="L")
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=education)
    pdf.ln(10)

    # Add Skills Section
    pdf.set_font("Arial", style='B', size=12)
    pdf.cell(200, 10, txt="SKILLS", ln=True, align="L")
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=f"- {skills.replace(',', '\n- ')}")
    pdf.ln(10)

    # Add Experience Section
    pdf.set_font("Arial", style='B', size=12)
    pdf.cell(200, 10, txt="EXPERIENCE", ln=True, align="L")
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=f"- {experience.replace(',', '\n- ')}")
    pdf.ln(10)

    # Save the PDF to a BytesIO object
    pdf_output = BytesIO()
    pdf_output.write(pdf.output(dest="S").encode("latin1"))
    pdf_output.seek(0)

    return pdf_output

# Function to export the resume to DOCX
def export_to_docx(name, job_role, resume_text, education, skills, experience, phone, email, linkedin, address):
    doc = Document()
    doc.add_heading(name.upper(), 0)
    doc.add_paragraph(f"Job Role: {job_role.upper()}")
    doc.add_paragraph(f"Phone: {phone} | Email: {email} | LinkedIn: {linkedin} | Address: {address}")
    
    doc.add_heading("PROFESSIONAL SUMMARY", level=1)
    doc.add_paragraph(resume_text)
    
    doc.add_heading("EDUCATION", level=1)
    doc.add_paragraph(education)
    
    doc.add_heading("SKILLS", level=1)
    doc.add_paragraph(skills.replace(',', '\n'))
    
    doc.add_heading("EXPERIENCE", level=1)
    doc.add_paragraph(experience.replace(',', '\n'))
    
    doc_output = BytesIO()
    doc.save(doc_output)
    doc_output.seek(0)

    return doc_output

# Streamlit interface
def main():
    st.title("ATS OPTIMIZED RESUME GENERATOR")

    # Collect user information
    st.header("Enter Your Information")
    name = st.text_input("Full Name", key="name")
    job_role = st.text_input("Job Role", key="job_role")
    education = st.text_input("Education (e.g., BSc in Computer Science)", key="education")
    skills = st.text_area("Skills (e.g., Python, JavaScript, etc.)", key="skills")
    experience = st.text_area("Experience (e.g., 2 years as Software Developer)", key="experience")
    
    # Add contact details inputs
    phone = st.text_input("Phone Number", key="phone")
    email = st.text_input("Your Email Address", key="email")
    linkedin = st.text_input("LinkedIn Profile", key="linkedin")
    address = st.text_input("Home Address", key="address")
    
    # Add Job Description Input (Optional)
    job_description = st.text_area("Job Description (Optional)", key="job_description")
    
    # File format and background color selection
    file_format = st.selectbox("Select File Format", ["PDF", "DOCX"], key="file_format")
    background_color = st.selectbox("Select Resume Background Color", ["white", "light grey", "blue"], key="background_color")

    # Button to generate resume
    if st.button("Generate Resume", key="generate_resume"):
        if name and job_role and education and skills and experience and phone and email and linkedin and address:
            resume_text = generate_resume(name, job_role, education, skills, experience, job_description)

            st.subheader("Generated Resume")
            st.write(resume_text)

            if file_format == "PDF":
                pdf_output = export_to_pdf(name, job_role, resume_text, education, skills, experience, phone, email, linkedin, address, background_color)
                st.download_button("Download Resume", pdf_output, file_name=f"{name}_Resume.pdf")
            else:
                docx_output = export_to_docx(name, job_role, resume_text, education, skills, experience, phone, email, linkedin, address)
                st.download_button("Download Resume", docx_output, file_name=f"{name}_Resume.docx")
        else:
            st.error("Please fill in all fields to generate a resume.")

if __name__ == "__main__":
    main()
